from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = "outputs/车路云一体化系统数据分类分级-车辆数据标准深度研究与修改建议.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_widths(table, widths_cm):
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def add_table(doc, headers, rows, widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, True)
        set_cell_shading(hdr[i], "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
    set_table_widths(table, widths_cm)
    doc.add_paragraph()
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_source(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.add_run(text)


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.59)
section.page_height = Cm(27.94)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)
section.header_distance = Cm(1.25)
section.footer_distance = Cm(1.25)

styles = doc.styles
for style_name in ["Normal", "List Bullet", "List Number"]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.1

for style_name, size, color, before, after in [
    ("Title", 22, "0B2545", 0, 8),
    ("Subtitle", 11, "555555", 0, 12),
    ("Heading 1", 16, "2E74B5", 16, 8),
    ("Heading 2", 13, "2E74B5", 12, 6),
    ("Heading 3", 12, "1F4D78", 8, 4),
]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("《车路云一体化系统数据分类分级：第一部分 车辆数据》\n深度研究与修改建议")
sub = doc.add_paragraph(style="Subtitle")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("研究对象：CSAE 标准文本 0210V5（PDF，2026-03-10 版本）")
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("形成日期：2026-05-21    适用场景：标准评审、征求意见反馈、编制组修改讨论")

doc.add_heading("一、研究结论", level=1)
doc.add_paragraph(
    "本标准的选题方向具有现实必要性。车路云一体化进入规模化试点后，车辆数据在车端、路侧、云控平台、监管部门和第三方服务机构之间流动，确实需要形成可执行的分类分级规则。0210V5 版本已经覆盖车辆数据清单、智能化网联化赋能、交互链路、时延、分级管控、交易流通和动态调整等内容，框架较完整。"
)
doc.add_paragraph(
    "但从标准成熟度看，当前文本仍存在四类核心问题：第一，标准对象边界还不够清晰，容易与自动驾驶数据记录系统、路侧数据、云端数据混淆；第二，第 5 章分类主线过多且未形成统一编码，INL1-INL3 更像车辆能力或应用场景分组，而不是车辆数据分类；第三，第 6 章分级规则与 GB/T 43697-2024 的一般数据、重要数据、核心数据框架尚未充分对齐；第四，部分安全保护、加密、交易要求过细或不现实，容易被专家认为偏管理制度或技术方案，而不是分类分级标准。"
)
doc.add_paragraph(
    "建议将本标准定位收敛为：面向车路云一体化场景的车辆数据分类维度、数据项清单、分级判定规则、典型应用映射和动态调整机制。文本结构宜采用“基础分类目录 + 多维标签 + 分级判定 + 典型数据项分级参考”的方式重构。"
)

doc.add_heading("二、建议重构后的标准主线", level=1)
add_number(doc, "先定义车辆数据基础分类目录，覆盖车辆基础标识与静态参数、车辆运行状态、控制与执行、智驾感知与环境、定位导航与轨迹、车路云交互、车载系统与通信、用户与座舱、诊断运维、能源动力、支付交易与服务等。")
add_number(doc, "再定义多维标签，包括交互链路、时效性、应用场景、主体角色、加工形态、个人信息属性、重要数据属性和商业秘密属性。")
add_number(doc, "最后依据 GB/T 43697-2024 的影响对象和危害程度进行分级，并给出 DL1-DL6 与一般数据、重要数据、核心数据之间的参考映射。")
add_number(doc, "附录 A 应从资料性清单升级为可落地的数据项级映射表，支撑企业数据资产盘点、接口设计、共享开放、交易流通和监管审查。")

doc.add_heading("三、逐章问题与修改建议", level=1)
rows = [
    ("范围", "1.1 将适用范围限定为 M 类和 N 类车辆配备的自动驾驶数据记录系统，边界过窄，且容易与 GB 44497-2024 混淆。", "改为覆盖车辆在车路、车云、车车、车人交互及相关服务中的采集、传输、存储、加工、共享、开放和交易活动；另写明涉及自动驾驶数据记录系统的数据项应符合 GB 44497-2024。"),
    ("术语", "“车辆核心数据”“车辆敏感数据”与《数据安全法》、GB/T 43697-2024、《个人信息保护法》中的核心数据、重要数据、敏感个人信息概念容易冲突。", "删除或重构自定义术语，统一引用核心数据、重要数据、一般数据、个人信息、敏感个人信息、商业秘密数据、安全关键数据等概念。"),
    ("总体原则", "原则性表述较多，如“分类多维”“科学实用性”“可行性验证”，但缺少可执行的方法和判据。", "将原则转化为可执行规则，例如数据项清单字段、分类标签、分级判定流程、升级/降级条件和主体适用规则。"),
    ("第 5 章分类", "INL1/INL2/INL3 存在 L0-L2、L1-L2 边界重叠；其本质是车辆能力或场景，不宜作为数据基础分类主线。", "将 INL 等级调整为应用标签或车辆能力标签；第 5 章主线改为车辆数据基础分类目录，再叠加交互链路、时效性、应用场景等标签。"),
    ("附录 A", "数据项粒度不一致，部分路端/云端下发数据混入车辆数据；缺少数据来源、个人信息属性、重要数据属性和建议等级。", "重构为数据项级清单，字段包括数据项、一级分类、二级分类、数据来源、交互链路、时效性、典型场景、个人信息属性、重要数据属性、建议等级、备注。"),
    ("第 6 章分级", "DL1-DL6 直接使用“核心级、关键级”等名称，可能被误解为等同国家核心数据或重要数据。", "明确 DL1-DL3 为一般数据内部细分，DL4-DL5 为重要数据候选或重要数据，DL6 为核心数据候选；最终识别以主管部门目录和监管要求为准。"),
    ("安全措施", "“量子加密”“明文传输”“丢包率 0.01%”等表述不适合作为分类分级标准中的通用要求。", "改为分级保护参考要求，并与等保、密码应用、车端实时性、V2X 通信安全要求衔接。"),
    ("数据交易", "6.2.6 仍沿用“分级响应、快速止损”表述，明显与交易主题不匹配。", "重写为禁止交易、严格受控使用、授权限制流通、可开放交易四类，并增加原始数据不出域、可用不可见、用途限定、禁止再识别等条件。"),
]
add_table(doc, ["章节", "主要问题", "修改建议"], rows, [2.2, 6.3, 7.6])

doc.add_heading("四、重点技术修改意见", level=1)
doc.add_heading("1. 车辆数据基础分类应先于场景分类", level=2)
doc.add_paragraph(
    "车路云一体化场景下，同一车辆数据项可能服务于多种场景。例如实时位置既可用于安全预警，也可用于交通治理、运维调度、保险服务和模型训练。若直接按 INL 或应用场景分类，会造成一个数据项在多个表中重复出现，难以形成统一目录。建议采用“一项数据、一个基础分类、多个应用标签”的方式。"
)
add_bullet(doc, "基础分类解决“这是什么数据”。")
add_bullet(doc, "交互链路标签解决“数据在哪里流动”。")
add_bullet(doc, "时效性标签解决“数据需要多快处理”。")
add_bullet(doc, "主体角色标签解决“谁可以在什么条件下使用”。")
add_bullet(doc, "合规属性标签解决“是否涉及个人信息、敏感个人信息、重要数据、商业秘密”。")

doc.add_heading("2. 车外数据需要专门条款", level=2)
doc.add_paragraph(
    "车路云系统中风险最高的数据往往不是单项 CAN 信号，而是车外图像、视频、点云、车牌、人脸、交通参与者轨迹，以及车辆与路侧、云端融合后形成的高精度时空数据。这类数据可能同时涉及个人信息、敏感个人信息、公共安全、国家安全和商业秘密。建议新增“车外数据处理要求”小节。"
)
add_bullet(doc, "原则上优先车端处理，默认不外传原始视频、图像、点云和可识别自然人的数据。")
add_bullet(doc, "确需外传时，应进行人脸、车牌、敏感区域过滤、空间精度降级、时间窗口控制和用途限定。")
add_bullet(doc, "批量、高精度、持续性车辆轨迹和交通参与者轨迹应作为重要数据候选进行评估。")
add_bullet(doc, "用于 AI 模型训练的数据集应区分原始样本、脱敏样本、匿名化样本、标签数据和统计数据。")

doc.add_heading("3. 分级示例需要按风险条件动态判定", level=2)
doc.add_paragraph(
    "当前文本中部分示例存在定级偏高或偏低问题。建议不要把某个数据项永久绑定某个等级，而是给出默认等级、升级条件和降级条件。这样更符合 GB/T 43697-2024 的风险导向和动态调整思路。"
)
rows = [
    ("公开车型参数", "DL1", "含未公开技术参数、供应链信息或商业秘密", "已由车企正式公开发布"),
    ("单车实时位置", "DL3", "连续轨迹、绑定用户身份、覆盖敏感区域", "粗粒度化、匿名化、不可回溯到单车或个人"),
    ("批量车辆轨迹", "DL4-DL5", "高精度、大规模、长时间、覆盖敏感区域或重要区域", "聚合统计、空间网格化、时间窗口扩大"),
    ("车外视频/图像", "DL4-DL5", "含人脸、车牌、敏感区域、事故现场或可识别自然人", "车端实时处理后不留存，或充分脱敏匿名化"),
    ("制动/转向/底盘控制指令", "DL5", "可远程控制或影响车辆行驶安全", "仅保留统计指标、仿真样本或不可执行记录"),
    ("密钥、证书私钥、安全芯片材料", "DL6 或禁止流转", "任何外传、复制、交易或非授权访问", "原则上不建议降级"),
    ("自动驾驶算法参数", "DL4-DL5", "涉及核心控制策略、漏洞利用路径或商业秘密", "抽象化指标、测试报告或不可复原结果"),
    ("匿名化交通流统计", "DL1-DL2", "可回溯到单车、个人或敏感区域规律", "不可复识别、不可还原、仅统计用途"),
]
add_table(doc, ["数据类型", "默认等级建议", "升级条件", "降级条件"], rows, [3.0, 2.5, 5.2, 5.4])

doc.add_heading("4. 时延分类应作为标签，不应与安全等级绑定", level=2)
doc.add_paragraph(
    "强实时、弱实时、非实时是资源调度和接口设计维度，不等同于安全等级。强实时数据可能只是公开交通信号状态，也可能是高风险控制指令；非实时数据也可能因大规模汇聚而成为重要数据。建议统一术语，并将阈值作为典型参考。"
)
add_bullet(doc, "强实时：典型端到端时延参考值 ≤100 ms 或 ≤200 ms，用于碰撞预警、协同控制、远程驾驶关键链路。")
add_bullet(doc, "准实时/弱实时：200 ms 至秒级或分钟级，用于状态监控、路径规划、交通调度。")
add_bullet(doc, "非实时：分钟级以上，用于统计分析、模型训练、审计、运维和历史追溯。")

doc.add_heading("五、建议写入征求意见反馈表的 12 条意见", level=1)
feedback = [
    "建议修改 1.1 范围，避免将本标准限定为自动驾驶数据记录系统，应覆盖车端、车路、车云、车车交互中的车辆数据全生命周期。",
    "建议删除或重构“车辆核心数据”“车辆敏感数据”术语，统一引用 GB/T 43697-2024、《数据安全法》《个人信息保护法》中的核心数据、重要数据、一般数据、敏感个人信息等概念。",
    "建议第 5 章先建立车辆数据基础分类目录，再基于交互链路、时效性、应用场景、主体角色、加工形态等设置扩展分类标签。",
    "建议将 INL1-INL3 从“数据分类主线”调整为“应用场景/车辆能力标签”，避免与驾驶自动化等级、网联化等级发生边界重叠。",
    "建议重构附录 A，形成数据项级清单，并增加数据来源、交互链路、时效性、个人信息属性、重要数据属性、建议等级等字段。",
    "建议第 6 章分级规则对齐 GB/T 43697-2024，明确 DL1-DL6 与一般数据、重要数据、核心数据之间的参考映射关系。",
    "建议修正部分分级示例，不宜将软件版本号、定位精度等单项数据直接定为 DL6，应结合规模、场景、精度、关联性和危害程度判断。",
    "建议对车外图像、视频、点云、人脸、车牌、交通参与者轨迹等设专门条款，明确车端处理、最小外传、脱敏匿名化和敏感区域过滤要求。",
    "建议统一“强实时、弱实时、非实时/低延迟”等术语，并将 100 ms、200 ms、分钟级等指标作为典型参考而非绝对要求。",
    "建议第 6.2 的权限、加密、审计、应急要求改为参考性保护措施，删除“量子加密”“明文传输”“固定丢包率”等不适合标准化或不现实的表述。",
    "建议重写数据交易分级规则，区分禁止交易、严格受控使用、授权限制流通、可开放交易四类，并引入“原始数据不出域、可用不可见、禁止再识别”等机制。",
    "建议补充动态调整的触发阈值，例如数据规模、定位精度、时间跨度、空间范围、是否覆盖敏感区域、是否可识别个人、是否跨主体融合等。",
]
for item in feedback:
    add_number(doc, item)

doc.add_heading("六、建议的附录 A 数据清单字段", level=1)
doc.add_paragraph(
    "附录 A 是本标准落地的关键。建议从现有资料性清单升级为可执行的数据项级清单，至少包含以下字段："
)
rows = [
    ("数据项", "例如实时经纬度、制动主缸压力、车外视频、车辆 VIN、功能订阅状态。"),
    ("一级/二级分类", "对应车辆数据基础分类目录，避免同一数据项多处重复。"),
    ("数据来源", "车身域、智驾域、控制域、T-Box、OBU、座舱、APP、云平台回传等。"),
    ("交互链路", "车云、车路、车车、车人、车内、平台间。"),
    ("时效性", "强实时、准实时/弱实时、非实时。"),
    ("典型场景", "安全预警、协同感知、协同决策、远程驾驶、交通治理、运维、模型训练、交易。"),
    ("合规属性", "个人信息、敏感个人信息、重要数据候选、商业秘密、安全关键数据、公开数据。"),
    ("建议等级", "DL1-DL6，并注明默认等级、升级条件、降级条件。"),
    ("处理建议", "本地处理、脱敏外传、授权使用、原始数据不出域、禁止交易等。"),
]
add_table(doc, ["字段", "说明"], rows, [3.2, 12.2])

doc.add_heading("七、参考资料", level=1)
sources = [
    "GB/T 43697-2024《数据安全技术 数据分类分级规则》，国家标准全文公开系统。",
    "《中华人民共和国数据安全法》，全国人民代表大会常务委员会，2021 年。",
    "《中华人民共和国个人信息保护法》，全国人民代表大会常务委员会，2021 年。",
    "《汽车数据安全管理若干规定（试行）》，国家互联网信息办公室等，2021 年。",
    "GB 44497-2024《智能网联汽车 自动驾驶数据记录系统》。",
    "GB/T 41871-2022《信息安全技术 汽车数据处理安全要求》。",
    "T/CSAE 313-2023《车路云一体化系统数据分类分级指南》。",
    "EDPB Guidelines 01/2020 on processing personal data in the context of connected vehicles and mobility related applications.",
    "European Commission, Guidance on vehicle data accompanying the EU Data Act.",
    "UNECE UN Regulation No.155, Cyber security and cyber security management system.",
    "ISO/SAE 21434 Road vehicles - Cybersecurity engineering.",
]
for s in sources:
    add_source(doc, s)

doc.add_section(WD_SECTION.NEW_PAGE)
doc.add_heading("附录：建议替换性表述示例", level=1)
doc.add_heading("1. 范围条款建议文本", level=2)
doc.add_paragraph(
    "本文件规定车路云一体化系统中车辆数据的分类维度、分级规则、典型数据项分级参考和动态调整要求。本文件适用于 M 类、N 类智能网联汽车在车路、车云、车车、车人交互及相关服务过程中产生、采集、传输、存储、加工、共享、开放和交易的车辆数据。涉及自动驾驶数据记录系统的数据项，应符合 GB 44497-2024 的要求。"
)
doc.add_heading("2. 分级映射建议文本", level=2)
doc.add_paragraph(
    "本文件采用 DL1-DL6 作为车路云一体化车辆数据安全等级参考。其中 DL1-DL3 可对应一般数据内部细分，DL4-DL5 可作为重要数据候选或重要数据管理参考，DL6 可作为核心数据候选管理参考。涉及重要数据、核心数据的最终识别，应以国家及行业主管部门发布的目录、规则和监管要求为准。"
)
doc.add_heading("3. 车外数据处理建议文本", level=2)
doc.add_paragraph(
    "车辆采集车外图像、视频、点云、车牌、人脸、交通参与者轨迹和道路环境数据时，应区分原始采集数据、实时安全处理数据、脱敏后共享数据和统计分析数据。原则上优先采用车端本地处理和最小必要外传；确需外传的，应进行去标识化、匿名化、敏感区域过滤、空间精度降级、时间窗口控制和用途限定。"
)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("车路云一体化系统数据分类分级：车辆数据标准研究与修改建议")

doc.save(OUT)
print(OUT)
