#!/usr/bin/env python3
"""Build front/back contract relationship candidates from workbook + Word contracts.

This script corrects an important modeling point:

- Excel detail sheets are NOT contracts. They are equipment/software/service
  list pools.
- Contract boundaries must come from signed front-sales/back-procurement
  contract documents.
- The tail section of workbook sheet `合同流` is used as the macro cash-flow
  skeleton.
- Front contracts and back contracts form many-to-many relationships through
  Tianan, and must be confirmed at item/amount/tax evidence level later.

The output is intentionally candidate-only. It must not be used for payment or
gross-margin calculation without human confirmation.
"""

from __future__ import annotations

import argparse
import json
import math
import re
import subprocess
import tempfile
from collections import Counter
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook


DEFAULT_WORKBOOK = Path(
    "/Users/tt2000/Documents/天安智联/0000无锡市车路云一体化项目/0.合同/合同管理/无锡车路云一体化采购价格测算表202500920.xlsx"
)
DEFAULT_DOCUMENT_ASSETS = Path("prototype/data/document-assets.json")
DEFAULT_OUT = Path("prototype/data/contract-relationships.json")
DEFAULT_DERIVED = Path("document_assets/derived/contract-relationships.json")

PARTY_ALIASES = {
    "无锡市车联网产业集团有限公司": "车联网",
    "无锡市车联网产业集团": "车联网",
    "车联网集团": "车联网",
    "车联网": "车联网",
    "江苏天安智联科技股份有限公司": "天安",
    "天安智联": "天安",
    "天安": "天安",
    "中国移动通信集团江苏有限公司无锡分公司": "移动",
    "中国移动": "移动",
    "江苏移动": "移动",
    "移动": "移动",
    "浪潮": "浪潮",
    "航天大为": "大为",
    "大为": "大为",
    "航天海特": "大为",
    "华通": "华通",
    "隆顺": "隆顺",
    "中信科智联": "中信科",
    "中信科": "中信科",
    "联通华盛": "联通华盛",
    "联通": "联通",
    "吉利": "吉利",
    "大华": "大华",
    "海康智联": "海康智联",
    "海康": "海康智联",
    "工业安装": "工业安装",
    "中邮建": "中邮建",
    "上研": "上研",
    "上电科": "上电科",
    "中通服网盈": "中通服网盈",
    "中通服盈科": "中通服网盈",
    "尚行": "尚行",
    "帆一": "尚行",
    "上汽帆一": "尚行",
    "合创": "合创",
    "万集": "万集",
    "车城": "车城",
    "车城智联": "车城",
    "四维图新": "四维图新",
    "信通院": "信通院",
    "交科所": "交科所",
    "金中天": "金中天",
    "金晓": "金晓",
    "中科创达": "中科创达",
    "华三": "华三",
    "华为": "华为",
    "回车": "回车",
    "东土": "东土",
    "恒信和安": "恒信和安",
    "博世": "博世",
    "创新中心": "创新中心",
    "无锡电信": "中电鸿信",
    "江苏电信": "中电鸿信",
}

NON_PARTY_TOKENS = {
    "R1",
    "R2",
    "R3",
    "R4",
    "RSU",
    "OBU",
    "GPU",
    "C-V2X",
    "服务器",
    "移动自有",
}

NON_PARTY_KEYWORDS = [
    "预付额度",
    "新国标",
    "可变标识",
    "低速无人",
    "机柜",
    "信号机软件",
    "POC安装",
    "公交硬件",
    "应用平台",
]

COMMON_FLOW_PARTIES = {"天安", "移动", "车联网"}

DEVICE_KEYWORDS = [
    "RSU", "信号机", "信号灯", "服务器", "GPU", "雷达", "激光雷达", "毫米波",
    "摄像", "相机", "边缘云", "云控", "高精地图", "地图", "数据质量", "C-V2X",
    "网关", "密码机", "公交", "车载", "OBU", "光纤", "物联网卡", "机柜",
    "施工", "安装", "运维", "软件", "平台", "算法", "三维实景",
]

ALLOWED_TAX_RATES = {"13%", "9%", "6%"}
ALLOWED_TAX_RATE_ORDER = ["13%", "9%", "6%"]

LEGACY_MACRO_FLOW_KEY = "移动-" + "吉利" + "-天安（大华+" + "希迪" + "+RSU)"
MACRO_FLOW_NAME_OVERRIDES = {
    LEGACY_MACRO_FLOW_KEY: "移动-四维图新-天安（大华+海康智联+RSU)",
}

PAYMENT_STAGE_TEMPLATES = [
    {
        "stageCode": "advance",
        "stageName": "预付款",
        "stageOrder": 10,
        "triggerCondition": "合同签署、预付款条款满足",
        "ownerAuditRequirement": "通常不要求设备审计，但需绑定后续设备级资金池",
        "upstreamRequirement": "背靠背条款下需上游预付款到账",
        "releaseRule": "按后向合同预付款条款释放；不得越过候选关系确认闸门",
        "partialReleaseAllowed": True,
    },
    {
        "stageCode": "delivery",
        "stageName": "到货款",
        "stageOrder": 20,
        "triggerCondition": "到货、入库、发票或送货签收证据齐备",
        "ownerAuditRequirement": "可按到货批次预审，最终受业主设备级审计确认金额约束",
        "upstreamRequirement": "对应设备到货阶段回款或设备级资金池可用",
        "releaseRule": "按到货数量和后向单价折算部分释放",
        "partialReleaseAllowed": True,
    },
    {
        "stageCode": "installation",
        "stageName": "安装款",
        "stageOrder": 30,
        "triggerCondition": "安装签证、安装数量、NodeID/点位绑定完成",
        "ownerAuditRequirement": "按安装点位进入业主审计或预审",
        "upstreamRequirement": "对应设备安装阶段回款或资金池可用",
        "releaseRule": "按实际安装数量折算释放，未绑定点位不得释放",
        "partialReleaseAllowed": True,
    },
    {
        "stageCode": "acceptance",
        "stageName": "验收款",
        "stageOrder": 40,
        "triggerCondition": "验收单、验收数量和问题闭环完成",
        "ownerAuditRequirement": "验收资料进入业主设备级审计",
        "upstreamRequirement": "对应验收阶段回款或资金池可用",
        "releaseRule": "按验收通过数量和审计前置规则释放",
        "partialReleaseAllowed": True,
    },
    {
        "stageCode": "owner_audit",
        "stageName": "审计款",
        "stageOrder": 50,
        "triggerCondition": "业主审计确认数量和金额",
        "ownerAuditRequirement": "必须取得正式设备级审计结果",
        "upstreamRequirement": "审计款到账或明确可归集到设备级资金池",
        "releaseRule": "以审计确认金额为基数，扣除核减后释放",
        "partialReleaseAllowed": True,
    },
    {
        "stageCode": "warranty",
        "stageName": "质保金",
        "stageOrder": 60,
        "triggerCondition": "质保期届满、无遗留问题或问题已闭环",
        "ownerAuditRequirement": "受最终审计确认金额和质保扣款约束",
        "upstreamRequirement": "质保/尾款到账或可用",
        "releaseRule": "扣除质量问题、审计核减和已付款后释放",
        "partialReleaseAllowed": True,
    },
]

DEVICE_CASHFLOW_SCHEMA = {
    "minimumUnit": "front_contract_item -> back_contract_item -> item_pool_row -> NodeID/site -> owner_audit -> payment_stage",
    "mustHaveFields": [
        "macroFlowId",
        "frontContractId",
        "backContractId",
        "itemName",
        "model",
        "quantity",
        "frontTaxRate",
        "backTaxRate",
        "frontAmountTaxIncluded",
        "backAmountTaxIncluded",
        "nodeId",
        "ownerAuditStatus",
        "ownerAuditConfirmedAmount",
        "upstreamReceiptReleasedAmount",
        "downstreamPaymentReleasedAmount",
        "overpaymentRiskStatus",
    ],
    "gateRule": "front/back candidates are financialUseAllowed=false until confirmed; confirmed links still require device-item split and owner-audit evidence before payment release.",
}


def clean(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).strip()
    return re.sub(r"\s+", " ", text)


def stable_id(prefix: str, *parts: Any) -> str:
    import hashlib

    raw = "|".join(clean(p) for p in parts)
    return f"{prefix}_{hashlib.sha1(raw.encode('utf-8')).hexdigest()[:16]}"


def money(value: Any) -> float | None:
    if value is None or value == "":
        return None
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        if math.isnan(value):
            return None
        return float(value)
    s = clean(value).replace(",", "").replace("￥", "")
    try:
        return float(s)
    except Exception:
        return None


def normalize_party(text: Any) -> str:
    s = clean(text)
    if not s:
        return ""
    normalized_token = s.strip("()（） ")
    normalized_token = re.sub(r"[)）]+$", "", normalized_token).strip()
    upper_token = normalized_token.upper()
    if upper_token in NON_PARTY_TOKENS or re.fullmatch(r"R\d+\s*等?", upper_token):
        return ""
    for key, val in PARTY_ALIASES.items():
        if key in s:
            return val
    if normalized_token in NON_PARTY_TOKENS:
        return ""
    if any(keyword in normalized_token for keyword in NON_PARTY_KEYWORDS):
        return ""
    # Ignore obvious notes/numbers.
    if re.fullmatch(r"[-+]?\d+(?:\.\d+)?", s) or len(s) > 40:
        return ""
    return s


def parties_in_text(text: str) -> list[str]:
    found = []
    for key, val in PARTY_ALIASES.items():
        if key in text and val not in found:
            found.append(val)
    return found


def split_chain(text: Any) -> list[str]:
    s = clean(text)
    if not s:
        return []
    # Treat brackets as separators. The previous escaped character-class form
    # missed Chinese parentheses, so values like `天安（华通` collapsed into
    # Tianan and lost the downstream party.
    s = re.sub(r"[【】\[\]（）()]", "-", s)
    parts = re.split(r"[-—→>+/、,，]+", s)
    out = []
    for part in parts:
        p = normalize_party(part)
        if p and p not in out:
            out.append(p)
    return out


def distinctive_flow_parties(flow: MacroFlow) -> set[str]:
    return set(flow.chain + flow.downstreamPartiesAfterTianan) - COMMON_FLOW_PARTIES


def upstream_parties_to_tianan(flow: MacroFlow) -> set[str]:
    if "天安" not in flow.chain:
        return set()
    return set(flow.chain[: flow.chain.index("天安")]) - COMMON_FLOW_PARTIES


def downstream_parties_from_tianan(flow: MacroFlow) -> set[str]:
    return set(flow.downstreamPartiesAfterTianan) - COMMON_FLOW_PARTIES


def normalize_macro_flow_name(name: str) -> str:
    return MACRO_FLOW_NAME_OVERRIDES.get(clean(name), clean(name))


def extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    parts.extend(p.text for p in doc.paragraphs if p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_doc_text(path: Path) -> str:
    """Extract legacy .doc text by converting through macOS textutil.

    python-docx cannot read binary .doc. In this local prototype environment,
    textutil can convert most .doc contracts to docx, after which the normal
    docx extractor is used. If conversion fails, return empty text and let
    the UI fall back to manual confirmation instead of inventing fields.
    """
    with tempfile.TemporaryDirectory(prefix="wuxi-contract-doc-") as tmp:
        out = Path(tmp) / f"{path.stem}.docx"
        try:
            subprocess.run(
                ["/usr/bin/textutil", "-convert", "docx", "-output", str(out), str(path)],
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                timeout=20,
            )
            if out.exists() and out.stat().st_size > 0:
                return extract_docx_text(out)
        except Exception:
            return ""
    return ""


def extract_amounts(text: str, limit: int = 20) -> list[float]:
    amounts = []
    for m in re.finditer(r"(?<![A-Za-z0-9])(\d{1,3}(?:,\d{3})+|\d+)(?:\.(\d+))?\s*(亿元|亿|万元|万|元)?", text):
        raw = m.group(0)
        n = float((m.group(1) + ("." + m.group(2) if m.group(2) else "")).replace(",", ""))
        unit = m.group(3) or ""
        if unit in ("亿元", "亿"):
            n *= 100000000
        elif unit in ("万元", "万"):
            n *= 10000
        elif not unit and n < 10000:
            continue
        if 1000 <= n <= 500000000 and n not in amounts:
            amounts.append(round(n, 2))
        if len(amounts) >= limit:
            break
    return amounts


def extract_tax_rates(text: str) -> list[str]:
    tax_context = "\n".join(line for line in text.splitlines() if any(k in line for k in ["税率", "税点", "含税", "增值税", "发票"]))
    out = []
    for val in re.findall(r"(\d{1,2}(?:\.\d+)?)\s*%", tax_context):
        s = normalize_tax_rate(f"{val}%")
        if s and s not in out:
            out.append(s)
    return out


def extract_contract_parties_from_text(text: str) -> dict[str, str]:
    result = {"partyAFullName": "", "partyBFullName": ""}
    if not text:
        return result
    patterns = [
        ("partyAFullName", r"甲方(?:（[^）]*）|\([^)]*\))?\s*[:：]\s*([^\n|；;，,]{4,80})"),
        ("partyBFullName", r"乙方(?:（[^）]*）|\([^)]*\))?\s*[:：]\s*([^\n|；;，,]{4,80})"),
    ]
    for key, pattern in patterns:
        match = re.search(pattern, text)
        if not match:
            continue
        value = clean(match.group(1))
        value = re.sub(r"\s+", "", value)
        value = re.sub(r"[。；;，,].*$", "", value).strip()
        if any(token in value for token in ["有限公司", "股份", "公司", "研究所", "集团"]):
            result[key] = value
    return result


def normalize_tax_rate(value: str) -> str:
    s = clean(value)
    if not s:
        return ""
    if "%" in s:
        m = re.search(r"(\d{1,2}(?:\.\d+)?)\s*%", s)
        if not m:
            return ""
        n = float(m.group(1))
    else:
        n = parse_number(s)
        if n is None:
            return ""
        if 0 < n < 1:
            n *= 100
    normalized = f"{n:g}%"
    return normalized if normalized in ALLOWED_TAX_RATES else ""


GENERIC_SHEET_KEYWORD_STOPWORDS = {
    "硬件",
    "软件",
    "服务",
    "购置",
    "硬件购置",
    "软件购置",
    "软件升级",
    "软件购置",
    "应用软件",
    "定制软件",
    "定制开发",
    "定制开发应用软件",
    "资源需求清单",
    "需求清单",
    "物理机版",
    "一阶段",
    "合同签署计划",
    "合同流",
    "分项报价汇总表",
    "硬件服务器",
    "更新",
    "CA",
    "密码",
    "POC",
    "安装费",
}


def strip_sheet_index(sheet_name: str) -> str:
    text = clean(sheet_name)
    text = re.sub(r"^[一二三四五六七八九十]+[.．、]\s*", "", text)
    text = re.sub(r"^\d+(?:\.\d+)*(?:R\d+)?\s*", "", text, flags=re.IGNORECASE)
    return text.strip()


def normalize_sheet_keyword(value: str) -> str:
    text = clean(value)
    text = re.sub(r"^[一二三四五六七八九十]+[.．、]\s*", "", text)
    text = re.sub(r"^\d+(?:\.\d+)*(?:R\d+)?\s*", "", text, flags=re.IGNORECASE)
    text = text.strip(" \t\r\n-—_，,、；;。.")
    text = re.sub(
        r"(?:硬件购置|软件购置|软件升级|应用软件|定制开发应用软件|定制软件|资源需求清单|需求清单|物理机版|一阶段)$",
        "",
        text,
    ).strip(" \t\r\n-—_，,、；;。.")
    return text


def workbook_sheet_keywords(workbook: Path) -> list[str]:
    wb = load_workbook(workbook, read_only=True, data_only=True)
    keywords: list[str] = []
    seen: set[str] = set()

    def add_keyword(value: str) -> None:
        keyword = normalize_sheet_keyword(value)
        if not keyword or keyword in GENERIC_SHEET_KEYWORD_STOPWORDS:
            return
        if len(keyword) < 2:
            return
        if "+" in keyword or "＋" in keyword:
            return
        if keyword.startswith("个"):
            return
        if keyword not in seen:
            seen.add(keyword)
            keywords.append(keyword)

    for sheet_name in wb.sheetnames:
        base = strip_sheet_index(sheet_name)
        if not base or base in GENERIC_SHEET_KEYWORD_STOPWORDS:
            continue

        for bracket_content in re.findall(r"【([^】]+)】", base):
            for token in re.split(r"[+＋/、,，\s]+", bracket_content):
                add_keyword(token)

        without_brackets = re.sub(r"【[^】]+】", "", base)
        add_keyword(without_brackets)
        for token in re.split(r"[+＋/、,，;；()（）\s]+", without_brackets):
            add_keyword(token)

        # Preserve meaningful prefixes before common suffix words. This turns
        # sheet names like “车联网视频平台硬件购置” into “车联网视频平台”.
        for token in list(keywords):
            add_keyword(token)

    return sorted(keywords, key=lambda item: (-len(item), item))


def keyword_hits(text: str, primary_keywords: list[str] | None = None) -> list[str]:
    lower = text.lower()
    out: list[str] = []
    for kw in primary_keywords or []:
        if kw.lower() in lower and kw not in out:
            out.append(kw)
    for kw in DEVICE_KEYWORDS:
        if kw.lower() in lower and kw not in out:
            out.append(kw)
    return out


@dataclass
class MacroFlow:
    id: str
    rowNumber: int
    packageName: str
    packageAmount: float | None
    chain: list[str]
    displayMode: str
    flowEdges: list[dict[str, Any]]
    displayText: str
    displayNote: str
    isTiananFrontSalesContract: bool
    tiananAmount: float | None
    downstreamPartiesAfterTianan: list[str]
    amounts: list[float]
    notes: str
    flowType: str


def extract_macro_flows(workbook: Path) -> list[MacroFlow]:
    wb = load_workbook(workbook, read_only=True, data_only=True)
    ws = wb["合同流"]
    flows: list[MacroFlow] = []
    current: dict[str, Any] | None = None
    for rn in range(138, ws.max_row + 1):
        package_name = clean(ws.cell(rn, 5).value)
        package_amount = money(ws.cell(rn, 6).value)
        row_values = [ws.cell(rn, c).value for c in range(1, 19)]
        row_text = " ".join(clean(v) for v in row_values if v is not None)
        if package_name and package_name != "天安合计":
            if current:
                flows.append(build_macro_flow(current))
            current = {
                "rowNumber": rn,
                "packageName": package_name,
                "packageAmount": package_amount,
                "rows": [row_values],
                "notes": [],
            }
        elif current:
            current["rows"].append(row_values)
        if current and row_text:
            current["notes"].append(row_text)
    if current:
        flows.append(build_macro_flow(current))
    return [flow for flow in flows if flow.isTiananFrontSalesContract]


def is_tianan_front_sales_macro_flow(package_name: str, chain: list[str]) -> bool:
    """Return true for macro rows that represent Tianan's front-sales contracts.

    The workbook contains both detailed package rows and a final front-contract
    skeleton. The contract-management page should show only rows whose standard
    name explicitly routes an upstream party to Tianan, such as
    `移动-浪潮-天安(...)` or `车联网-天安(...)`.
    """

    if "天安" not in package_name or "天安" not in chain:
        return False
    if chain.index("天安") <= 0:
        return False
    return bool(re.search(r"[-—→>]", package_name))


def build_macro_flow(raw: dict[str, Any]) -> MacroFlow:
    parties = []
    amounts = []
    notes = " ".join(raw.get("notes", []))
    package_name = normalize_macro_flow_name(raw["packageName"])
    notes = notes.replace(raw["packageName"], package_name)
    for row in raw["rows"]:
        for c in [4, 6, 8, 10, 12, 14, 16]:
            if c < len(row):
                p = normalize_party(row[c])
                if p and p not in parties:
                    parties.append(p)
        for c in [5, 7, 9, 11, 13, 15, 17]:
            if c < len(row):
                n = money(row[c])
                if n is not None:
                    amounts.append(round(n, 2))
    chain = split_chain(package_name)
    for p in parties:
        if p not in chain:
            chain.append(p)
    after_tianan = []
    if "天安" in chain:
        after_tianan = chain[chain.index("天安") + 1 :]
    is_tianan_front_sales = is_tianan_front_sales_macro_flow(package_name, chain)
    flow_type = (
        "tianan_front_sales"
        if is_tianan_front_sales
        else "tianan_direct"
        if chain and chain[0] == "天安"
        else "mobile_to_tianan"
        if "移动" in chain and "天安" in chain
        else "other"
    )
    display_mode, flow_edges, display_text, display_note = build_macro_flow_display(chain, after_tianan, flow_type)
    tianan_amount = raw["packageAmount"] if "天安" in chain else None
    return MacroFlow(
        id=stable_id("macro_flow", raw["rowNumber"], raw["packageName"], raw["packageAmount"]),
        rowNumber=raw["rowNumber"],
        packageName=package_name,
        packageAmount=raw["packageAmount"],
        chain=chain,
        displayMode=display_mode,
        flowEdges=flow_edges,
        displayText=display_text,
        displayNote=display_note,
        isTiananFrontSalesContract=is_tianan_front_sales,
        tiananAmount=tianan_amount,
        downstreamPartiesAfterTianan=after_tianan,
        amounts=amounts,
        notes=notes[:1200],
        flowType=flow_type,
    )


def build_macro_flow_display(
    chain: list[str],
    after_tianan: list[str],
    flow_type: str,
) -> tuple[str, list[dict[str, Any]], str, str]:
    """Build a display-safe macro-flow representation.

    `chain` is retained as an extraction clue. It must not be blindly rendered
    as a serial subcontracting path because Excel often puts several Tianan
    downstream suppliers in adjacent rows/columns. In that case the business
    meaning is Tianan -> A, Tianan -> B, Tianan -> C in parallel.
    """

    if not chain:
        return "unknown", [], "-", "未识别到主体链路，需人工复核。"

    if flow_type == "tianan_front_sales" and "天安" in chain:
        tianan_index = chain.index("天安")
        upstream = chain[: tianan_index + 1]
        edges = [
            {
                "from": upstream[idx],
                "to": upstream[idx + 1],
                "amount": None,
                "edgeType": "front_sales_to_tianan",
                "evidenceStatus": "parsed_from_macro_flow_name",
            }
            for idx in range(len(upstream) - 1)
        ]
        return (
            "front_sales",
            edges,
            " → ".join(upstream),
            "仅展示天安前向销售资金入口；括号或后续主体只作为后向设备级匹配候选，不代表宏观流骨架中的串行链。",
        )

    if "天安" in chain and after_tianan:
        tianan_index = chain.index("天安")
        upstream = chain[: tianan_index + 1]
        edges: list[dict[str, Any]] = []
        for idx in range(len(upstream) - 1):
            edges.append(
                {
                    "from": upstream[idx],
                    "to": upstream[idx + 1],
                    "amount": None,
                    "edgeType": "serial_upstream_to_tianan",
                    "evidenceStatus": "parsed_from_macro_flow_name",
                }
            )
        for party in after_tianan:
            edges.append(
                {
                    "from": "天安",
                    "to": party,
                    "amount": None,
                    "edgeType": "parallel_downstream_from_tianan",
                    "evidenceStatus": "parsed_from_excel_downstream_party",
                }
            )

        if tianan_index == 0:
            display_text = "；".join(f"天安 → {party}" for party in after_tianan)
        else:
            upstream_text = " → ".join(upstream)
            display_text = f"{upstream_text}；" + "；".join(f"天安 → {party}" for party in after_tianan)

        return (
            "parallel" if flow_type == "tianan_direct" else "mixed_parallel",
            edges,
            display_text,
            "并联拆分：天安分别流向多个后向主体，不能按相邻主体顺序理解为串行链；边金额需按合同附件/设备清单复核。",
        )

    if len(chain) == 1:
        return "standalone", [], chain[0], "单节点宏观流，需结合原始行和合同文本复核。"

    edges = [
        {
            "from": chain[idx],
            "to": chain[idx + 1],
            "amount": None,
            "edgeType": "serial",
            "evidenceStatus": "parsed_from_macro_flow_name",
        }
        for idx in range(len(chain) - 1)
    ]
    return "serial", edges, " → ".join(chain), "按名称暂按串行展示，仍需合同文本与合同附件确认。"


@dataclass
class ContractDoc:
    id: str
    fileName: str
    sourcePath: str
    direction: str
    counterparty: str
    parties: list[str]
    partyAFullName: str
    partyBFullName: str
    amounts: list[float]
    taxRates: list[str]
    keywords: list[str]
    textLength: int
    backToBackLikely: bool
    evidenceLines: list[str]


@dataclass
class ContractItem:
    id: str
    contractId: str
    contractName: str
    direction: str
    counterparty: str
    sourcePath: str
    tableIndex: int
    rowNumber: int
    itemNo: str
    itemName: str
    detailName: str
    specModel: str
    quantity: float | None
    unit: str
    unitPriceTaxIncluded: float | None
    amountTaxIncluded: float | None
    amountTaxExcluded: float | None
    taxRate: str
    rawCells: list[str]
    evidenceLevel: str


def record_priority(record: dict) -> tuple[int, str]:
    source_path = str(record.get("sourcePath", ""))
    # Prefer source-of-truth contract folders over manual import copies when the
    # same file hash appears in both places. Manual imports should supplement,
    # not duplicate, already indexed contracts.
    if "contract_manual_imports" in source_path:
        return (1, source_path)
    return (0, source_path)


def dedupe_document_asset_records(records: list[dict]) -> list[dict]:
    by_hash: dict[str, dict] = {}
    no_hash: list[dict] = []
    for record in records:
        file_hash = record.get("sha256")
        if not file_hash:
            no_hash.append(record)
            continue
        current = by_hash.get(file_hash)
        if current is None or record_priority(record) < record_priority(current):
            by_hash[file_hash] = record
    return sorted([*by_hash.values(), *no_hash], key=lambda item: str(item.get("sourcePath", "")))


def load_contract_exclusions(exclude_file: Path | None) -> list[dict[str, Any]]:
    if not exclude_file or not exclude_file.exists():
        return []
    try:
        payload = json.loads(exclude_file.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return []
    return payload.get("exclusions", []) if isinstance(payload, dict) else []


def record_is_excluded(record: dict[str, Any], exclusions: list[dict[str, Any]]) -> bool:
    if not exclusions:
        return False
    source_path = str(record.get("sourcePath", "")).strip()
    file_hash = str(record.get("sha256", "")).strip()
    for exclusion in exclusions:
        excluded_path = str(exclusion.get("sourcePath", "")).strip()
        excluded_hash = str(exclusion.get("sha256", "")).strip()
        if excluded_path and excluded_path == source_path:
            return True
        if excluded_hash and excluded_hash == file_hash:
            return True
    return False


def extract_contract_docs(document_assets: Path, primary_keywords: list[str] | None = None, exclusions: list[dict[str, Any]] | None = None) -> list[ContractDoc]:
    data = json.loads(document_assets.read_text(encoding="utf-8"))
    docs = []
    for record in dedupe_document_asset_records(data.get("records", [])):
        if record_is_excluded(record, exclusions or []):
            continue
        path = Path(record["sourcePath"])
        ext = path.suffix.lower()
        if ext not in {".docx", ".doc", ".pdf"}:
            continue
        text = ""
        if ext == ".docx":
            try:
                text = extract_docx_text(path)
            except Exception:
                text = ""
        elif ext == ".doc":
            text = extract_doc_text(path)
        elif ext == ".pdf":
            # PDF contracts are accepted into the contract register, but this
            # prototype does not parse them. They remain file-level records for
            # manual confirmation and OS/default-app reading.
            text = ""
        fields = record.get("extractedFields", {})
        filename = record.get("fileName", path.name)
        base_text = f"{filename}\n{fields.get('contractName','')}\n{fields.get('partyA','')}\n{fields.get('partyB','')}\n{fields.get('counterparty','')}\n{text}"
        parties = parties_in_text(base_text)
        for key in [fields.get("partyA"), fields.get("partyB"), fields.get("counterparty")]:
            p = normalize_party(key)
            if p and p not in parties:
                parties.append(p)
        amount_candidates = []
        if fields.get("amountCny"):
            amount_candidates.append(float(fields["amountCny"]))
        for amount in extract_amounts(text):
            if amount not in amount_candidates:
                amount_candidates.append(amount)
        taxes = extract_tax_rates(text)
        party_full_names = extract_contract_parties_from_text(text)
        kws = keyword_hits(base_text, primary_keywords)
        lines = []
        for line in text.splitlines():
            compact = clean(line)
            if compact and (any(k in compact for k in kws[:8]) or any(t in compact for t in taxes)):
                lines.append(compact[:260])
            if len(lines) >= 8:
                break
        docs.append(
            ContractDoc(
                id=stable_id("contract_doc", filename, record.get("sourcePath")),
                fileName=filename,
                sourcePath=record.get("sourcePath", ""),
                direction=fields.get("contractDirection") or record.get("documentSubcategory", ""),
                counterparty=normalize_party(fields.get("counterparty", "")),
                parties=parties,
                partyAFullName=party_full_names.get("partyAFullName", ""),
                partyBFullName=party_full_names.get("partyBFullName", ""),
                amounts=amount_candidates[:20],
                taxRates=taxes,
                keywords=kws,
                textLength=len(text),
                backToBackLikely=any(k in text for k in ["甲方客户", "业主", "同等比例", "背靠背", "回款"]),
                evidenceLines=lines,
            )
        )
    return docs


def parse_number(value: str) -> float | None:
    s = clean(value).replace(",", "").replace("，", "")
    if not s or s in {"-", "—", "/"}:
        return None
    m = re.search(r"-?\d+(?:\.\d+)?", s)
    if not m:
        return None
    try:
        return float(m.group(0))
    except Exception:
        return None


def parse_tax_rate(value: str) -> str:
    return normalize_tax_rate(value)


def normalize_header(text: str) -> str:
    return re.sub(r"[\s\n\r\t（）()：:；;、,，/\\]+", "", clean(text)).lower()


ITEM_HEADER_RULES: list[tuple[str, list[str]]] = [
    ("itemNo", ["序号", "no.", "no", "编号"]),
    ("itemName", ["采购品名", "分项名称", "项目", "类别", "品名", "设备名称", "服务名称", "子系统名称", "系统名称"]),
    ("detailName", ["采购明细", "明细", "一级模块", "二级模块", "模块", "内容"]),
    ("specModel", ["品牌及型号", "规格型号", "型号", "性能指标参数", "参数", "品牌", "服务器版本"]),
    ("quantity", ["数量", "工作量", "工作量人月", "数据数量"]),
    ("unit", ["单位"]),
    ("unitPriceTaxIncluded", ["含税单价", "综合单价", "单价"]),
    ("amountTaxIncluded", ["含税小计", "含税金额", "含税价小计", "金额元", "金额（元）", "金额"]),
    ("amountTaxExcluded", ["不含税价小计", "不含税金额", "未税金额"]),
    ("taxRate", ["税率", "税点"]),
]


def header_field(cell_text: str) -> str:
    h = normalize_header(cell_text)
    if not h:
        return ""
    for field, keys in ITEM_HEADER_RULES:
        if any(normalize_header(key) in h for key in keys):
            return field
    return ""


def find_item_header_row(table: Any) -> tuple[int, dict[str, int]] | None:
    max_scan = min(8, len(table.rows))
    best: tuple[int, dict[str, int], int] | None = None
    for row_idx in range(max_scan):
        mapping: dict[str, int] = {}
        for col_idx, cell in enumerate(table.rows[row_idx].cells):
            field = header_field(cell.text)
            if field and field not in mapping:
                mapping[field] = col_idx
        score = len(mapping)
        if "itemName" in mapping or "detailName" in mapping:
            score += 2
        if "amountTaxIncluded" in mapping or "amountTaxExcluded" in mapping:
            score += 2
        if "taxRate" in mapping:
            score += 1
        if score >= 5 and (best is None or score > best[2]):
            best = (row_idx, mapping, score)
    if not best:
        return None
    return best[0], best[1]


def cell_value(cells: list[str], header_map: dict[str, int], field: str) -> str:
    idx = header_map.get(field)
    if idx is None or idx >= len(cells):
        return ""
    return clean(cells[idx])


def useful_item_text(*parts: str) -> str:
    joined = " ".join(clean(p) for p in parts if clean(p))
    joined = re.sub(r"\b(?:合计|总计|小计|项目总价款|服务税率\d+%?)\b", " ", joined, flags=re.IGNORECASE)
    return clean(joined)


def extract_contract_items(docs: list[ContractDoc]) -> list[ContractItem]:
    items: list[ContractItem] = []
    for doc in docs:
        path = Path(doc.sourcePath)
        if path.suffix.lower() != ".docx":
            continue
        try:
            word_doc = Document(str(path))
        except Exception:
            continue
        for table_idx, table in enumerate(word_doc.tables):
            header = find_item_header_row(table)
            if not header:
                continue
            header_row, header_map = header
            for row_idx in range(header_row + 1, len(table.rows)):
                cells = [clean(cell.text).replace("\n", " / ") for cell in table.rows[row_idx].cells]
                if not any(cells):
                    continue
                item_no = cell_value(cells, header_map, "itemNo")
                item_name = cell_value(cells, header_map, "itemName")
                detail_name = cell_value(cells, header_map, "detailName")
                spec_model = cell_value(cells, header_map, "specModel")
                tax_rate = parse_tax_rate(cell_value(cells, header_map, "taxRate"))
                amount_tax_included = money(cell_value(cells, header_map, "amountTaxIncluded"))
                amount_tax_excluded = money(cell_value(cells, header_map, "amountTaxExcluded"))
                quantity = parse_number(cell_value(cells, header_map, "quantity"))
                unit_price = money(cell_value(cells, header_map, "unitPriceTaxIncluded"))
                unit = cell_value(cells, header_map, "unit")
                item_text = useful_item_text(item_name, detail_name, spec_model)
                if not item_text:
                    continue
                if any(skip in item_text for skip in ["合作涉及数据情况", "安全要求", "用户上网日志", "数据分级", "总计", "合计"]):
                    continue
                has_financial_signal = any(v is not None for v in [amount_tax_included, amount_tax_excluded, unit_price]) or bool(tax_rate)
                has_quantity_signal = quantity is not None and (unit or unit_price is not None)
                if not has_financial_signal and not has_quantity_signal:
                    continue
                if amount_tax_included is None and unit_price is not None and quantity is not None:
                    amount_tax_included = round(unit_price * quantity, 2)
                item_id = stable_id(
                    "contract_item",
                    doc.id,
                    table_idx,
                    row_idx + 1,
                    item_no,
                    item_name,
                    detail_name,
                    spec_model,
                    amount_tax_included,
                    amount_tax_excluded,
                    tax_rate,
                )
                items.append(
                    ContractItem(
                        id=item_id,
                        contractId=doc.id,
                        contractName=doc.fileName,
                        direction=doc.direction,
                        counterparty=doc.counterparty,
                        sourcePath=doc.sourcePath,
                        tableIndex=table_idx,
                        rowNumber=row_idx + 1,
                        itemNo=item_no,
                        itemName=item_name or detail_name or spec_model,
                        detailName=detail_name,
                        specModel=spec_model,
                        quantity=quantity,
                        unit=unit,
                        unitPriceTaxIncluded=unit_price,
                        amountTaxIncluded=amount_tax_included,
                        amountTaxExcluded=amount_tax_excluded,
                        taxRate=tax_rate,
                        rawCells=cells[:16],
                        evidenceLevel="word_table_row",
                    )
                )
    return items


def amount_close(a: float, b: float) -> bool:
    if not a or not b:
        return False
    tolerance = max(5000, min(a, b) * 0.03)
    return abs(a - b) <= tolerance


MATCH_STOPWORDS = {
    "项目",
    "服务",
    "设备",
    "采购",
    "购置",
    "定制",
    "开发",
    "应用",
    "软件",
    "硬件",
    "系统",
    "平台",
    "合同",
    "车路云",
    "一体化",
    "无锡",
    "元",
}


def item_match_tokens(item: ContractItem) -> set[str]:
    text = " ".join([item.itemName, item.detailName, item.specModel])
    upper = text.upper()
    tokens: set[str] = set()
    for kw in DEVICE_KEYWORDS:
        if kw.upper() in upper:
            tokens.add(kw.upper())
    for token in re.findall(r"[A-Za-z][A-Za-z0-9_.+-]{1,}|[0-9]+(?:\.[0-9]+)?|[\u4e00-\u9fff]{2,}", text):
        normalized = token.upper() if re.search(r"[A-Za-z]", token) else token
        if normalized in MATCH_STOPWORDS:
            continue
        if len(normalized) >= 2:
            tokens.add(normalized)
    return tokens


def amount_ratio(front_amount: float | None, back_amount: float | None) -> float | None:
    if not front_amount or not back_amount:
        return None
    if front_amount <= 0 or back_amount <= 0:
        return None
    return round(back_amount / front_amount, 4)


def amount_relation_text(front_amount: float | None, back_amount: float | None) -> str:
    ratio = amount_ratio(front_amount, back_amount)
    if ratio is None:
        return "金额待确认"
    if amount_close(front_amount or 0, back_amount or 0):
        return "金额接近"
    if 0 < ratio <= 1:
        return f"后向约为前向 {ratio:.1%}"
    return f"后向高于前向 {ratio:.1%}"


def build_device_item_match_candidates(
    front_items: list[ContractItem],
    back_items: list[ContractItem],
    contract_candidates: list[dict[str, Any]],
    limit: int = 1200,
) -> list[dict[str, Any]]:
    contract_context: dict[tuple[str, str], dict[str, Any]] = {}
    for candidate in contract_candidates:
        contract_context[(candidate["frontContractId"], candidate["backContractId"])] = candidate

    back_tokens = {item.id: item_match_tokens(item) for item in back_items}
    front_tokens = {item.id: item_match_tokens(item) for item in front_items}
    candidates: list[dict[str, Any]] = []

    for front in front_items:
        ftokens = front_tokens[front.id]
        if not ftokens:
            continue
        local_matches: list[dict[str, Any]] = []
        for back in back_items:
            btokens = back_tokens[back.id]
            overlap = sorted(ftokens & btokens)
            score = 0
            reasons = []

            context = contract_context.get((front.contractId, back.contractId))
            if context:
                score += 5
                reasons.append("合同级上下文命中同一前向/后向候选")

            if overlap:
                score += min(12, len(overlap) * 2)
                reasons.append("设备/服务词重合:" + ",".join(overlap[:8]))

            if front.taxRate and back.taxRate:
                if front.taxRate == back.taxRate:
                    score += 4
                    reasons.append(f"税率一致:{front.taxRate}")
                else:
                    score -= 3
                    reasons.append(f"税率不一致:{front.taxRate}/{back.taxRate}")

            if front.quantity is not None and back.quantity is not None:
                if abs(front.quantity - back.quantity) <= max(1, abs(front.quantity) * 0.03):
                    score += 5
                    reasons.append(f"数量接近:{front.quantity:g}≈{back.quantity:g}")
                elif 0 < back.quantity <= front.quantity:
                    score += 2
                    reasons.append(f"后向数量为前向子集:{back.quantity:g}/{front.quantity:g}")

            if front.unit and back.unit and front.unit == back.unit:
                score += 1
                reasons.append(f"单位一致:{front.unit}")

            if front.amountTaxIncluded and back.amountTaxIncluded:
                ratio = amount_ratio(front.amountTaxIncluded, back.amountTaxIncluded)
                if amount_close(front.amountTaxIncluded, back.amountTaxIncluded):
                    score += 6
                    reasons.append("含税金额接近")
                elif ratio and 0 < ratio <= 1:
                    score += 2
                    reasons.append("后向含税金额不超过前向")
                elif ratio and ratio > 1.15:
                    score -= 2
                    reasons.append("后向金额明显高于前向，需核查")

            if score < 8:
                continue

            macro_ids = context.get("sharedMacroFlowIds", []) if context else []
            local_matches.append(
                {
                    "id": stable_id("device_item_match", front.id, back.id),
                    "frontItemId": front.id,
                    "backItemId": back.id,
                    "frontContractId": front.contractId,
                    "frontContractName": front.contractName,
                    "backContractId": back.contractId,
                    "backContractName": back.contractName,
                    "frontItemName": front.itemName,
                    "frontDetailName": front.detailName,
                    "frontSpecModel": front.specModel,
                    "backItemName": back.itemName,
                    "backDetailName": back.detailName,
                    "backSpecModel": back.specModel,
                    "frontQuantity": front.quantity,
                    "backQuantity": back.quantity,
                    "unit": front.unit or back.unit,
                    "frontTaxRate": front.taxRate,
                    "backTaxRate": back.taxRate,
                    "taxRateMatch": bool(front.taxRate and back.taxRate and front.taxRate == back.taxRate),
                    "frontAmountTaxIncluded": front.amountTaxIncluded,
                    "backAmountTaxIncluded": back.amountTaxIncluded,
                    "amountRatio": amount_ratio(front.amountTaxIncluded, back.amountTaxIncluded),
                    "amountRelation": amount_relation_text(front.amountTaxIncluded, back.amountTaxIncluded),
                    "tokenOverlap": overlap[:16],
                    "sharedMacroFlowIds": macro_ids,
                    "score": score,
                    "confidence": "high" if score >= 20 else "medium" if score >= 13 else "low",
                    "status": "candidate",
                    "financialUseAllowed": False,
                    "requiresManualConfirmation": True,
                    "confirmationTarget": "front_contract_item_to_back_contract_item",
                    "reasons": reasons,
                    "importantNote": "这是设备/服务明细项级候选。必须人工逐条确认名称、规格型号、数量、单价、金额、税率、合同附件来源和审计口径后，才能形成设备级合同管理关联关系。",
                }
            )
        local_matches.sort(key=lambda x: (-x["score"], x["backContractName"], x["backItemName"]))
        candidates.extend(local_matches[:8])

    candidates.sort(key=lambda x: (-x["score"], x["frontContractName"], x["frontItemName"], x["backContractName"]))
    return candidates[:limit]


def match_contract_to_flows(doc: ContractDoc, flows: list[MacroFlow]) -> list[dict[str, Any]]:
    matches = []
    doc_parties = set(doc.parties)
    doc_keywords = set(doc.keywords)
    counterparty = normalize_party(doc.counterparty)
    front_party_anchor = counterparty if counterparty and counterparty not in COMMON_FLOW_PARTIES else ""
    if doc.direction == "front_sales" and not front_party_anchor:
        for part in Path(doc.sourcePath).parts:
            path_party = normalize_party(part)
            if path_party and path_party not in COMMON_FLOW_PARTIES:
                upstream_hit = any(path_party in upstream_parties_to_tianan(flow) for flow in flows)
                if upstream_hit:
                    front_party_anchor = path_party
                    break
    if doc.direction == "front_sales" and not front_party_anchor:
        candidate_upstream_parties = sorted(doc_parties - COMMON_FLOW_PARTIES)
        if len(candidate_upstream_parties) == 1:
            front_party_anchor = candidate_upstream_parties[0]
    for flow in flows:
        score = 0
        reasons = []
        distinctive_overlap = doc_parties & distinctive_flow_parties(flow)
        upstream_overlap = doc_parties & upstream_parties_to_tianan(flow)
        downstream_overlap = doc_parties & downstream_parties_from_tianan(flow)
        public_overlap = doc_parties & set(flow.chain + flow.downstreamPartiesAfterTianan) & COMMON_FLOW_PARTIES

        if upstream_overlap:
            score += len(upstream_overlap) * 4
            reasons.append(f"前向链路主体重合:{','.join(sorted(upstream_overlap))}")
        if downstream_overlap:
            score += len(downstream_overlap) * 4
            reasons.append(f"天安后向主体重合:{','.join(sorted(downstream_overlap))}")
        if distinctive_overlap and not (upstream_overlap or downstream_overlap):
            score += len(distinctive_overlap) * 2
            reasons.append(f"区分主体重合:{','.join(sorted(distinctive_overlap))}")
        if public_overlap and (upstream_overlap or downstream_overlap):
            score += 1
            reasons.append(f"公共主体辅助:{','.join(sorted(public_overlap))}")

        # Public parties such as Tianan/Mobile/owner are too common in this
        # project. They cannot by themselves make a flow match, otherwise every
        # contract appears to match every front-sales entry.
        if doc.direction == "front_sales" and not (upstream_overlap or distinctive_overlap):
            pass
        if doc.direction == "back_procurement" and not downstream_overlap:
            pass

        if doc.direction == "front_sales" and not (upstream_overlap or distinctive_overlap):
            # Keep only strong amount/keyword evidence below for very sparse
            # docs; no score is added for merely containing Tianan.
            pass
        if doc.direction == "front_sales" and front_party_anchor:
            # A Tianan front-sales contract's macro-flow ownership is primarily
            # determined by the upstream party that brings funds into Tianan.
            # Device/service keywords such as 信号机、服务器、平台 are item-level
            # evidence and must not move a 尚行/帆一 contract into 浪潮/上研等
            # unrelated funding-entry chains.
            if front_party_anchor not in upstream_parties_to_tianan(flow):
                continue
        if doc.direction == "back_procurement" and not downstream_overlap:
            # A back-procurement contract should match a front-sales entry only
            # when its Tianan downstream party is in that entry's downstream
            # clue set. Amount/keyword evidence can still add limited support.
            pass
        amount_hits = []
        for a in doc.amounts[:10]:
            for b in [flow.packageAmount or 0, *(flow.amounts[:20])]:
                if amount_close(a, b):
                    amount_hits.append((a, b))
                    break
        if amount_hits:
            score += min(6, len(amount_hits) * 2)
            reasons.append("金额接近:" + ",".join(f"{int(a)}≈{int(b)}" for a, b in amount_hits[:3]))
        keyword_overlap = [kw for kw in doc_keywords if kw in flow.packageName or kw in flow.notes]
        if keyword_overlap:
            score += min(3, len(keyword_overlap))
            reasons.append("设备/服务关键词:" + ",".join(keyword_overlap[:5]))
        if doc.direction == "front_sales":
            has_required_anchor = bool(upstream_overlap or distinctive_overlap or amount_hits or keyword_overlap)
        elif doc.direction == "back_procurement":
            has_required_anchor = bool(downstream_overlap)
        else:
            has_required_anchor = bool(distinctive_overlap)
        if score >= 4 and has_required_anchor:
            matches.append(
                {
                    "macroFlowId": flow.id,
                    "macroFlowPackage": flow.packageName,
                    "macroFlowRow": flow.rowNumber,
                    "score": score,
                    "reasons": reasons,
                    "confidence": "high" if score >= 10 else "medium" if score >= 6 else "low",
                }
            )
    return sorted(matches, key=lambda x: (-x["score"], x["macroFlowRow"]))[:8]


def build_front_back_candidates(docs: list[ContractDoc], flows: list[MacroFlow]) -> list[dict[str, Any]]:
    front_docs = [d for d in docs if d.direction == "front_sales"]
    back_docs = [d for d in docs if d.direction == "back_procurement"]
    flow_matches = {d.id: match_contract_to_flows(d, flows) for d in docs}
    candidates = []
    for front in front_docs:
        front_flow_ids = {m["macroFlowId"] for m in flow_matches.get(front.id, [])}
        for back in back_docs:
            back_flow_ids = {m["macroFlowId"] for m in flow_matches.get(back.id, [])}
            shared_flow = sorted(front_flow_ids & back_flow_ids)
            party_overlap = sorted((set(front.parties) & set(back.parties)) - {"天安"})
            keyword_overlap = sorted(set(front.keywords) & set(back.keywords))
            tax_overlap = sorted(set(front.taxRates) & set(back.taxRates))
            score = 0
            reasons = []
            if shared_flow:
                score += 7
                reasons.append("命中同一宏观合同流")
            if "天安" in front.parties and "天安" in back.parties:
                score += 3
                reasons.append("均通过天安")
            if party_overlap:
                score += min(4, len(party_overlap) * 2)
                reasons.append("共同相对方/下游主体:" + ",".join(party_overlap[:4]))
            if keyword_overlap:
                score += min(5, len(keyword_overlap))
                reasons.append("设备/服务关键词重合:" + ",".join(keyword_overlap[:6]))
            if tax_overlap:
                score += 1
                reasons.append("税率重合:" + ",".join(tax_overlap))
            # Avoid flooding with weak pairs.
            if score < 8:
                continue
            candidates.append(
                {
                    "id": stable_id("front_back_candidate", front.id, back.id, ",".join(shared_flow)),
                    "frontContractId": front.id,
                    "frontContractName": front.fileName,
                    "backContractId": back.id,
                    "backContractName": back.fileName,
                    "sharedMacroFlowIds": shared_flow,
                    "partyOverlap": party_overlap,
                    "keywordOverlap": keyword_overlap[:12],
                    "taxRateOverlap": tax_overlap,
                    "score": score,
                    "confidence": "high" if score >= 14 else "medium",
                    "status": "candidate",
                    "financialUseAllowed": False,
                    "reasons": reasons,
                    "importantNote": "这是前后向合同多对多候选关系，不是设备级最终确认。必须继续比对合同附件清单、金额、税率、数量和履约资料。",
                }
            )
    return sorted(candidates, key=lambda x: (-x["score"], x["frontContractName"], x["backContractName"]))


def find_consortium_master_contract(docs: list[ContractDoc]) -> ContractDoc | None:
    """Find the owner-to-consortium master contract.

    In this project, Tianan's apparent front-sales contracts are allocations or
    pass-through entries derived from the owner/consortium master agreement.
    The master contract should therefore be modeled as the upstream funding
    source of all Tianan front-sales macro-flow entries, not as a normal peer
    front-sales contract line.
    """
    candidates: list[tuple[int, ContractDoc]] = []
    for doc in docs:
        if doc.direction != "front_sales":
            continue
        text = f"{doc.fileName} {doc.sourcePath}"
        score = 0
        if "一阶段合同协议书" in text:
            score += 8
        if "车联网集团" in text or "车联网" in doc.parties:
            score += 4
        if "联合体" in " ".join(doc.evidenceLines) or {"移动", "天安"} <= set(doc.parties):
            score += 3
        if doc.amounts and max(doc.amounts[:5]) >= 100_000_000:
            score += 3
        if score:
            candidates.append((score, doc))
    if not candidates:
        return None
    candidates.sort(key=lambda item: (-item[0], item[1].fileName))
    return candidates[0][1]


def build_master_contract_record(doc: ContractDoc | None) -> dict[str, Any] | None:
    if not doc:
        return None
    amount = next((amount for amount in doc.amounts if amount >= 100_000_000), doc.amounts[0] if doc.amounts else None)
    return {
        "id": doc.id,
        "fileName": doc.fileName,
        "sourcePath": doc.sourcePath,
        "amountCny": amount,
        "contractRole": "owner_to_consortium_master_contract",
        "fundingRole": "所有天安侧前向资金入口的总合同来源；不是单条资金入口金额。",
        "parties": doc.parties,
    }


def macro_flow_record(flow: MacroFlow, master_contract: dict[str, Any] | None) -> dict[str, Any]:
    record = asdict(flow)
    record["amountRole"] = "tianan_front_sales_allocation"
    record["amountEvidence"] = "workbook_sheet_合同流"
    record["amountNote"] = "该金额是从联合体总合同分解/流转到天安侧的前向资金入口金额，不是业主联合体总合同金额。"
    if master_contract:
        record["sourceMasterContractId"] = master_contract["id"]
        record["sourceMasterContractName"] = master_contract["fileName"]
        record["sourceMasterContractAmountCny"] = master_contract["amountCny"]
    return record


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--workbook", type=Path, default=DEFAULT_WORKBOOK)
    parser.add_argument("--document-assets", type=Path, default=DEFAULT_DOCUMENT_ASSETS)
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT)
    parser.add_argument("--derived", type=Path, default=DEFAULT_DERIVED)
    parser.add_argument("--exclude-file", type=Path, default=None)
    args = parser.parse_args()

    macro_flows = extract_macro_flows(args.workbook)
    sheet_keywords = workbook_sheet_keywords(args.workbook)
    exclusions = load_contract_exclusions(args.exclude_file)
    docs = extract_contract_docs(args.document_assets, sheet_keywords, exclusions)
    contract_items = extract_contract_items(docs)
    front_items = [item for item in contract_items if item.direction == "front_sales"]
    back_items = [item for item in contract_items if item.direction == "back_procurement"]
    doc_flow_matches = [
        {"contractId": d.id, "contractName": d.fileName, "direction": d.direction, "matches": match_contract_to_flows(d, macro_flows)}
        for d in docs
    ]
    candidates = build_front_back_candidates(docs, macro_flows)
    device_item_candidates = build_device_item_match_candidates(front_items, back_items, candidates)
    master_contract = build_master_contract_record(find_consortium_master_contract(docs))
    payload = {
        "schemaVersion": 1,
        "source": {
            "workbook": str(args.workbook),
            "documentAssets": str(args.document_assets),
            "primaryKeywordSource": "workbook_sheet_names",
        },
        "rules": {
            "allowedTaxRates": ALLOWED_TAX_RATE_ORDER,
            "taxRateRule": "合同关系重建只接受 13%、9%、6% 三类税率；其他百分比只能保留在原始证据文本中，不进入结构化税率字段、匹配字段或统计字段。",
        },
        "summary": {
            "macroFlowCount": len(macro_flows),
            "contractDocumentCount": len(docs),
            "frontContractCount": sum(1 for d in docs if d.direction == "front_sales"),
            "backContractCount": sum(1 for d in docs if d.direction == "back_procurement"),
            "masterContractAmountCny": master_contract.get("amountCny") if master_contract else None,
            "masterContractFileName": master_contract.get("fileName") if master_contract else "",
            "frontContractItemCount": len(front_items),
            "backContractItemCount": len(back_items),
            "frontBackCandidateCount": len(candidates),
            "deviceItemMatchCandidateCount": len(device_item_candidates),
            "taxRateCounts": dict(Counter(t for d in docs for t in d.taxRates)),
            "primaryKeywordCount": len(sheet_keywords),
            "excludedContractCount": len(exclusions),
            "note": "Macro-flow skeleton only keeps Tianan front-sales contract rows. Word contract tables provide front/back item candidates. Financial confirmation must happen at front-contract-item to back-contract-item level.",
        },
        "primaryKeywords": sheet_keywords,
        "masterContracts": [master_contract] if master_contract else [],
        "macroFlows": [macro_flow_record(f, master_contract) for f in macro_flows],
        "contracts": [asdict(d) for d in docs],
        "frontContractItems": [asdict(item) for item in front_items],
        "backContractItems": [asdict(item) for item in back_items],
        "contractToMacroFlowMatches": doc_flow_matches,
        "frontBackRelationshipCandidates": candidates,
        "deviceItemMatchCandidates": device_item_candidates,
        "deviceCashflowSchema": DEVICE_CASHFLOW_SCHEMA,
        "ownerAuditTrackingModel": {
            "minimumAuditUnit": "cashflow_unit_id + node_id/site_id + item_name + audit_stage",
            "statuses": ["not_submitted", "submitted", "under_review", "confirmed", "partially_confirmed", "deducted", "evidence_required"],
            "rule": "Upstream receipt and downstream payment release must use owner audit confirmed quantity/amount, not only contract-level amount.",
        },
        "paymentStageTemplates": PAYMENT_STAGE_TEMPLATES,
    }
    for path in [args.out, args.derived]:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(payload["summary"], ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
